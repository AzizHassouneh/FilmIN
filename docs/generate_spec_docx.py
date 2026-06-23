#!/usr/bin/env python3
"""Generate FilmIN — Product Requirements & Scenarios Spec as a .docx for Google Docs."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---- palette ----
NAVY   = RGBColor(0x15, 0x23, 0x3b)
INK    = RGBColor(0x1f, 0x27, 0x33)
SLATE  = RGBColor(0x56, 0x60, 0x74)
AMBER  = RGBColor(0xd9, 0x8c, 0x2b)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
CARD   = RGBColor(0xf5, 0xf7, 0xfb)
LINE   = RGBColor(0xd8, 0xde, 0xe8)
AMBER_LIGHT = RGBColor(0xfb, 0xf1, 0xe2)

OUT = "/Users/azizhassouneh/filmin/docs/FilmIN-Product-Spec.docx"

# ---- helpers ----
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(rgb))
    tcPr.append(shd)

def set_cell_borders(cell, color="D8DEE8", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb

def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
    return p

def add_body(doc, text, bold_spans=None, italic=False, color=INK, size=10, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.italic = italic
    return p

def add_bullet(doc, text, color=INK, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def section_rule(doc, title):
    """Amber rule + heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D98C2B")
    pBdr.append(bottom)
    pPr.append(pBdr)
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY
    return p

def kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = AMBER

def persona_table(doc, cards):
    """2-column grid of persona cards, 2 per row."""
    rows = [cards[i:i+2] for i in range(0, len(cards), 2)]
    for pair in rows:
        while len(pair) < 2:
            pair.append(None)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.autofit = False
        tbl.columns[0].width = Inches(3.3)
        tbl.columns[1].width = Inches(3.3)
        row = tbl.rows[0]
        row.height = None
        for ci, card in enumerate(pair):
            cell = row.cells[ci]
            if card is None:
                set_cell_bg(cell, WHITE)
                set_cell_borders(cell, "FFFFFF")
                continue
            set_cell_bg(cell, CARD)
            set_cell_borders(cell, "E1E7F0")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Name line
            np = cell.paragraphs[0]
            np.paragraph_format.space_after = Pt(2)
            nr = np.add_run(f"{card['name']}  —  ")
            nr.font.bold = True
            nr.font.size = Pt(9.5)
            nr.font.color.rgb = NAVY
            nr2 = np.add_run(card["role"])
            nr2.font.size = Pt(9.5)
            nr2.font.color.rgb = SLATE
            for key in ("goals", "frustration", "filmin"):
            	lp = cell.add_paragraph()
            	lp.paragraph_format.space_after = Pt(1)
            	rb = lp.add_run({"goals": "Goals: ", "frustration": "Frustration: ", "filmin": "FilmIN gives them: "}[key])
            	rb.font.bold = True
            	rb.font.size = Pt(8.5)
            	rb.font.color.rgb = INK
            	rv = lp.add_run(card[key])
            	rv.font.size = Pt(8.5)
            	rv.font.color.rgb = INK
        doc.add_paragraph()  # spacer between card rows

def make_table(doc, headers, rows, col_widths, header_bg=NAVY, alt_bg=CARD):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = False
    for i, w in enumerate(col_widths):
        for cell in [tbl.columns[i].cells[0]]:
            cell.width = Inches(w)
    # header row
    hrow = tbl.rows[0]
    for ci, h in enumerate(headers):
        cell = hrow.cells[ci]
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = WHITE
    # data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = WHITE if ri % 2 == 0 else alt_bg
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            r.font.color.rgb = INK
    return tbl

def set_col_width(tbl, col_idx, width_inches):
    for cell in tbl.columns[col_idx].cells:
        cell.width = Inches(width_inches)

# ---- document ----
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin   = Inches(0.95)
        section.right_margin  = Inches(0.95)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = INK

    # ---- COVER ----
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("FilmIN")
    tr.font.bold = True
    tr.font.size = Pt(48)
    tr.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Product Requirements & Scenarios Spec")
    sr.font.size = Pt(16)
    sr.font.color.rgb = SLATE

    doc.add_paragraph()

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tgr = tag_p.add_run(
        "The open film & TV catalog of IMDb, fused with a LinkedIn-style\n"
        "professional network — built for the film industry."
    )
    tgr.font.size = Pt(12)
    tgr.font.italic = True
    tgr.font.color.rgb = INK

    doc.add_paragraph()

    badge_tbl = doc.add_table(rows=1, cols=1)
    badge_tbl.autofit = False
    badge_tbl.columns[0].width = Inches(5)
    bc = badge_tbl.rows[0].cells[0]
    set_cell_bg(bc, NAVY)
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run("FREE FOREVER  ·  FUNDED BY ADS, NOT PAYWALLS")
    br.font.bold = True
    br.font.size = Pt(11)
    br.font.color.rgb = WHITE
    badge_tbl._tbl.getparent()  # ensure in body
    # center the table
    badge_tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run("Working codename: FilmIN  ·  Draft v0.1  ·  June 21, 2026")
    mr.font.size = Pt(8.5)
    mr.font.color.rgb = SLATE

    doc.add_page_break()

    # ---- 1. Vision ----
    section_rule(doc, "1. Vision")
    add_body(doc,
        "FilmIN is the professional home for everyone who makes film — and the best place for fans "
        "to explore it. One trustworthy catalog of titles and credits, fused with a real professional "
        "network: owned profiles, connections, a feed, discovery, and hiring. Free for everyone, "
        "funded by ads, never by gatekeeping people from their own work.")

    # ---- 2. Value Prop ----
    section_rule(doc, "2. Value Proposition — “Why FilmIN”")
    vp = [
        ("Free forever for users", "page curation, headshot/photo uploads, credits management, discovery analytics, search, and contact tools that IMDbPro charges for are all free here."),
        ("Owned professional identity", "you control your page, your reel, your story."),
        ("Real industry networking", "connect, follow, endorse, post, get discovered, find collaborators and work."),
        ("A catalog both sides trust", "accurate, TMDB-powered titles and credits that fans and pros use."),
        ("Funded by ads, not paywalls", "monetization (later) never blocks a pro from managing their page."),
    ]
    for bold, rest in vp:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(bold + " — ")
        rb.font.bold = True; rb.font.size = Pt(9.5); rb.font.color.rgb = INK
        rr = p.add_run(rest)
        rr.font.size = Pt(9.5); rr.font.color.rgb = INK

    # ---- 3. Problem ----
    section_rule(doc, "3. The Problem Today")
    problems = [
        ("IMDb", "is resented; its moat is brand authority, not product quality. IMDbPro paywalls self-curation — you pay $150/yr to upload your own headshot."),
        ("LinkedIn", "is generic — no titles, no credits graph, no reels, no fan layer; not film-literate."),
        ("Letterboxd", "is fans-only — no professional identity or hiring layer."),
        ("Fragmentation", "pros juggle IMDb + LinkedIn + a personal site + a Vimeo reel + spreadsheets. FilmIN unifies all of it, free."),
    ]
    for bold, rest in problems:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(bold + " "); rb.font.bold = True; rb.font.size = Pt(9.5); rb.font.color.rgb = INK
        rr = p.add_run(rest); rr.font.size = Pt(9.5); rr.font.color.rgb = INK

    # ---- 4. Personas ----
    section_rule(doc, "4. Personas")
    add_body(doc,
        "Grouped as Creators (the pros), Industry enablers (who hire, represent, or program), "
        "Audience (fans & press), and Advertisers (future). Each card: goals · frustrations · what FilmIN gives them.")

    kicker(doc, "CREATORS — FILM PROFESSIONALS")
    creators = [
        {"name": "Maya",  "role": "Aspiring Actor",            "goals": "Visibility, polished headshots + reel, get cast.", "frustration": "Paying IMDbPro just to upload a headshot.", "filmin": "Claim page, curate, upload photos free; be discovered by casting."},
        {"name": "Sam",   "role": "Cinematographer / DP",      "goals": "Accurate credits, peer connections, signal availability.", "frustration": "Credits scattered; no film-aware network.", "filmin": "Verified filmography, 'open to work,' worked-with network, posts."},
        {"name": "Lia",   "role": "Choreographer",             "goals": "Be searchable in an under-catalogued role; show a reel.", "frustration": "Specialty roles poorly represented on IMDb.", "filmin": "First-class support for below-the-line roles + showreel + discovery."},
        {"name": "Dev",   "role": "Writer / Screenwriter",     "goals": "Showcase produced + spec work; reach directors/producers.", "frustration": "No good way to present a body of writing.", "filmin": "Credits + projects + networking + endorsements."},
        {"name": "Noor",  "role": "Director",                  "goals": "Showcase filmography, assemble crews, announce projects.", "frustration": "Hiring + portfolio live in different places.", "filmin": "Profile, people search, worked-with graph, feed."},
        {"name": "Alex",  "role": "Producer",                  "goals": "Discover/assemble talent & crew, manage a slate, network.", "frustration": "Sourcing trusted crew is slow and gated.", "filmin": "People search, shortlists, worked-with graph."},
        {"name": "Theo",  "role": "Film Student / Emerging",   "goals": "Build a profile with no IMDb credits yet; first break.", "frustration": "IMDb won't list student/indie work; invisible.", "filmin": "Add your own titles + credits; build identity from scratch."},
        {"name": "Priya", "role": "Costume Designer",          "goals": "Correct credit attribution; peer recognition.", "frustration": "Miscredited / missing on titles.", "filmin": "Claim, fix/add credits, endorsements."},
    ]
    persona_table(doc, creators)

    kicker(doc, "INDUSTRY ENABLERS")
    enablers = [
        {"name": "Carmen",   "role": "Casting Director",       "goals": "Find & shortlist talent by role/location/skills; contact them.", "frustration": "IMDbPro gates search & contact behind a fee.", "filmin": "Rich people search + filters + shortlists + free direct contact."},
        {"name": "Jordan",   "role": "Agent / Manager",        "goals": "Manage and promote a roster of clients.", "frustration": "No film-native tools to promote talent.", "filmin": "Manage client pages, promote, network."},
        {"name": "Studio",   "role": "Production Co. / Recruiter", "goals": "Source crews, scout, post opportunities (future).", "frustration": "Crewing-up relies on word of mouth.", "filmin": "Search + worked-with graph + company presence (future)."},
        {"name": "Festival", "role": "Festival Programmer",    "goals": "Discover filmmakers; partnerships and promotion.", "frustration": "Hard to find/track emerging filmmakers.", "filmin": "Discovery + partnership surfaces (a marketing channel)."},
    ]
    persona_table(doc, enablers)

    kicker(doc, "AUDIENCE  ·  (FUTURE) ADVERTISERS")
    audience = [
        {"name": "Joe",    "role": "General Fan",              "goals": "'Who was in that?', release dates, fast answers.", "frustration": "Cluttered, ad-heavy lookups.", "filmin": "Fast, clean catalog; follow talent."},
        {"name": "Quinn",  "role": "Cinephile (Letterboxd crowd)", "goals": "Track films; follow talent & critics; rate/review.", "frustration": "Fan tools and pro data live apart.", "filmin": "Catalog + follows + watchlists/reviews (future)."},
        {"name": "Rae",    "role": "Critic / Journalist",      "goals": "Publish reviews/articles; build authority; get cited.", "frustration": "Work gets cited to IMDb, not to them.", "filmin": "Critic pages + reviews; cite FilmIN over IMDb (future)."},
        {"name": "Brandr", "role": "Advertiser / Brand (future)", "goals": "Reach a film-intent audience.", "frustration": "Generic ad networks lack film context.", "filmin": "Ad-friendly surfaces designed without harming UX."},
    ]
    persona_table(doc, audience)

    doc.add_page_break()

    # ---- 5. Segments ----
    section_rule(doc, "5. User Segments")
    make_table(doc,
        ["Segment", "Share", "How FilmIN wins them"],
        [
            ["SUBSET ONE\nGeneral public",                       "~90%",  "Catalog lookups. Won by speed, accuracy, SEO, brand."],
            ["SUBSET TWO\nWorking / semi-pro\n[THE WEDGE]",      "~9%",   "Want free self-curation. Easiest to win; they evangelize to peers (grassroots growth)."],
            ["SUBSET THREE\nEstablished pros & companies",       "<1%",   "Previously IMDbPro payers. Pay nothing here; won by a better product + network effect."],
        ],
        col_widths=[1.8, 0.7, 4.1],
    )

    doc.add_paragraph()

    # ---- 6. Capabilities ----
    section_rule(doc, "6. Capabilities & Feature Areas")
    add_body(doc, "Tag [MVP] = in the first thin vertical slice; [Later] = roadmap.", color=SLATE, size=8.5)
    caps = [
        ("[MVP] Catalog & Discovery", "title pages (poster, year, synopsis), cast/crew, title search, trending rows, and a Discover browse page (role chips, trending, new releases, people to follow)."),
        ("[MVP] Professional Identity", "claim a page, profile (bio, roles, location, links), free headshot/photo upload. [Later] skills/tags, reel gallery."),
        ("[MVP] Credits & Filmography", "profile-title credits. [Later] add missing titles/credits (student/indie), correction requests."),
        ("[MVP] Professional Network", "follow, worked-with (derived from shared titles), and 'people you may know' suggestions. [Later] mutual connections, endorsements, 'open to work.'"),
        ("[MVP] Feed & Posts", "a personalized network feed on Home (posts + worked-with activity from your network, ranked by recency and proximity, with an extended-network tier), posts/updates, likes. [Later] comments."),
        ("[Later] Hiring & Discovery", "people search by role/location/skills, shortlists, direct contact (all free)."),
        ("[MVP] Fan Engagement", "follow talent. [Later] watchlists, ratings/reviews."),
        ("[Later] Critic & Editorial", "critic pages, reviews/articles."),
        ("[Later] Notifications", ""),
        ("[Later] Trust & Safety", "verification badges (claimed/verified pro), reporting/moderation."),
        ("[Later] Ads & Monetization", "ad-friendly layout considered now; implementation later."),
    ]
    for tag, desc in caps:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(tag); rb.font.bold = True; rb.font.size = Pt(9.5); rb.font.color.rgb = AMBER
        if desc:
            rr = p.add_run(" — " + desc); rr.font.size = Pt(9.5); rr.font.color.rgb = INK

    # ---- 7. Graph ----
    section_rule(doc, "7. The FilmIN Graph (Relationships)")
    add_body(doc, "The product is a graph connecting people, their work, and the audience:")
    graph_items = [
        ("Person —[Credit: role/character/job]— Title", "the spine of the catalog."),
        ("Person —[Worked-with]— Person", "derived automatically from shared titles."),
        ("User —[Owns/Claims]— Profile", "a profile is an unclaimed stub until a user claims it."),
        ("User —[Follows]→ Profile", "(future) Person ↔ Person mutual connection + endorsement."),
        ("Person —[Authors]— Post —[Reaction]— User", ""),
        ("Fan —[Follows]→ Talent", "(future) Fan —[Watchlist/Rating]— Title."),
    ]
    for bold, rest in graph_items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(bold); rb.font.bold = True; rb.font.size = Pt(9.5); rb.font.color.rgb = NAVY
        if rest:
            rr = p.add_run(" — " + rest); rr.font.size = Pt(9.5); rr.font.color.rgb = INK

    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run(
        "Node types: Profile, Title, User (account), Post, Fan\n"
        "Key edges: credit, claims/follows, authors, worked-with (shared titles)\n"
        "Profile ↔ Profile: worked-with is derived, not stored — shared credits create the link."
    )
    nr.font.size = Pt(8.5); nr.font.color.rgb = SLATE; nr.font.italic = True

    doc.add_page_break()

    # ---- 8. Scenarios ----
    section_rule(doc, "8. Scenarios (End-to-End Use Cases)")
    add_body(doc,
        "Each scenario: persona · trigger · steps · outcome. [MVP] = supported by the first thin slice.",
        color=SLATE, size=8.5)
    scenarios = [
        ("[MVP]  S1 — The free headshot (Maya).",
         "Signs up, searches her imported page, clicks 'This is me' to claim it, edits bio/roles, "
         "and uploads headshots for free. Value: the headline IMDbPro-killer."),
        ("[Later]  S2 — Casting a choreographer (Carmen → Lia).",
         "Searches 'Choreographer, LA, musical theatre,' shortlists Lia, views her reel, and "
         "contacts her — all free. Value: hiring without paywalls."),
        ("[Later]  S3 — Building from zero (Theo).",
         "With no IMDb credits, creates a profile, adds his student films as titles + credits, "
         "and connects with classmates. Value: emerging talent gets an identity IMDb won't give."),
        ("[MVP]  S4 — Open to work (Sam → Noor).",
         "Sam marks 'open to work' and posts a wrap announcement; Noor (in Sam's network) sees it "
         "on her Home feed and connects. Posts from friends-of-friends surface in an extended-network "
         "tier, and 'people you may know' suggests collaborators. Value: the network creates work."),
        ("[MVP]  S5 — Fan lookup (Joe).",
         "Searches a film; the title page shows cast/crew linked to profiles; he follows an actor. "
         "Value: fast, credible catalog; fan→talent follow graph."),
        ("[Later]  S6 — Fixing a credit (Priya).",
         "Finds a missing/mis-attributed credit on a title and requests a correction; the credit "
         "is fixed. Value: accuracy owned by the people who did the work."),
        ("[Later]  S7 — Assembling a crew (Alex).",
         "Uses people search and the worked-with graph to assemble a trusted crew."),
        ("[Later]  S8 — Critic citations (Rae).",
         "Publishes reviews on a critic page; outlets cite FilmIN over IMDb."),
        ("[MVP]  S9 — Following talent + worked-with (Quinn).",
         "Browses people on Discover, follows talent, and explores who they've worked with via shared titles."),
    ]
    for head, body_text in scenarios:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        rh = p.add_run(head); rh.font.bold = True; rh.font.size = Pt(10); rh.font.color.rgb = NAVY
        pb = doc.add_paragraph()
        pb.paragraph_format.left_indent = Inches(0.2)
        pb.paragraph_format.space_after = Pt(4)
        rb = pb.add_run(body_text); rb.font.size = Pt(9.5); rb.font.color.rgb = INK

    # ---- 9. Differentiators ----
    section_rule(doc, "9. Differentiators (vs the Field)")
    tbl = make_table(doc,
        ["Capability", "IMDb", "IMDbPro", "LinkedIn", "Letterboxd", "FilmIN"],
        [
            ["Open film/TV catalog",                         "Yes",       "Yes",       "No",        "Fan",       "Yes"],
            ["Self-curate page / upload headshots",          "No",        "Paid",      "Generic",   "No",        "Yes (free)"],
            ["Film-aware credits graph",                     "View only", "Yes",       "No",        "No",        "Yes"],
            ["Professional networking & hiring",             "No",        "Partial/paid", "Generic", "No",       "Yes (free)"],
            ["Fan engagement layer",                         "Yes",       "No",        "No",        "Yes",       "Yes"],
            ["Cost to professionals",                        "Free (view)","$150/yr",  "Free/paid", "Free",      "$0"],
        ],
        col_widths=[2.3, 0.78, 0.85, 0.85, 0.92, 0.92],
    )
    # highlight FilmIN column
    for ri in range(1, 7):
        cell = tbl.rows[ri].cells[5]
        set_cell_bg(cell, AMBER_LIGHT)

    doc.add_paragraph()

    # ---- 10. Principles ----
    section_rule(doc, "10. Principles (Non-Functional)")
    for item in [
        "Free core forever; the user owns their data and identity.",
        "Catalog accuracy; accessibility; fast and mobile-friendly.",
        "SEO for fan traffic (SUBSET ONE growth).",
        "Privacy-respecting; legal (TMDB attribution; no IMDb data or scraping).",
        "Ad-friendly, but never UX-hostile.",
    ]:
        add_bullet(doc, item)

    # ---- 11. Metrics ----
    section_rule(doc, "11. Success Metrics")
    for item in [
        "Claimed profiles; headshots uploaded; % of credits on claimed pages.",
        "Connections/follows created; weekly active users (pros vs fans).",
        "Catalog coverage; casting/people searches (future); grassroots referral signups.",
    ]:
        add_bullet(doc, item)

    # ---- 12. Monetization ----
    section_rule(doc, "12. Monetization (Future, Post-Traction)")
    add_body(doc,
        "Ads · sponsored discovery placements · festival/casting partnerships · product placement. "
        "Hard rule: never paywall a professional's own page or core networking. "
        "Outcome per the brief: scale ad revenue, or acquisition.")

    # ---- 13. Glossary ----
    section_rule(doc, "13. Glossary")
    glossary = [
        ("Title", "a film or TV work."),
        ("Credit", "a person's role on a title."),
        ("Profile", "a person's page; an unclaimed stub until claimed by its owner."),
        ("Claim", "proving/asserting a page is yours."),
        ("Follow (one-way) vs Connection (mutual, future)", ""),
        ("Worked-with", "a link between people derived from shared titles."),
        ("Reel", "a showcase video."),
        ("Endorsement", "a peer vouch (future)."),
    ]
    for term, defn in glossary:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(term); rb.font.bold = True; rb.font.size = Pt(9.5); rb.font.color.rgb = NAVY
        if defn:
            rr = p.add_run(" — " + defn); rr.font.size = Pt(9.5); rr.font.color.rgb = INK

    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "This product uses the TMDB API but is not endorsed or certified by TMDB. "
        "No IMDb data is used or scraped."
    )
    fr.font.size = Pt(8); fr.font.color.rgb = SLATE; fr.font.italic = True

    doc.save(OUT)
    print(f"WROTE {OUT}")

if __name__ == "__main__":
    build()
